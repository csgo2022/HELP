"""
生成点餐小程序项目文档 .docx
"""
from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn
import os

doc = Document()

# ============ 样式设置 ============
style = doc.styles['Normal']
font = style.font
font.name = '宋体'
font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

# 设置默认段落间距
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# 标题样式
for i in range(1, 5):
    hs = doc.styles[f'Heading {i}']
    hs.font.name = '黑体'
    hs.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    hs.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    if i == 1:
        hs.font.size = Pt(22)
    elif i == 2:
        hs.font.size = Pt(18)
    elif i == 3:
        hs.font.size = Pt(15)
    else:
        hs.font.size = Pt(13)

def add_table(doc, headers, rows, col_widths=None):
    """添加表格"""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # 表头
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(11)
    # 数据行
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            cell = table.rows[ri + 1].cells[ci]
            cell.text = str(val)
            for p in cell.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(11)
    return table

# ============ 封面 ============
for _ in range(6):
    doc.add_paragraph()

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('点餐小程序（顾客端）')
run.font.size = Pt(28)
run.bold = True
run.font.color.rgb = RGBColor(0xFF, 0x6B, 0x35)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('项目设计文档')
run.font.size = Pt(24)
run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()
doc.add_paragraph()

info_items = [
    '后端技术：Node.js + Express + MySQL',
    '前端平台：微信小程序',
    '文档版本：v1.0',
    '文档日期：2026年5月15日',
]
for item in info_items:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

doc.add_page_break()

# ============ 目录页 ============
doc.add_heading('目  录', level=1)
toc_items = [
    ('1', '绪论', 3),
    ('2', '需求分析', 4),
    ('3', '总体设计', 7),
    ('4', '数据库设计', 9),
    ('5', '接口设计', 11),
    ('6', '详细设计', 16),
    ('7', 'API 接口汇总', 20),
]
for num, title, page in toc_items:
    p = doc.add_paragraph()
    run = p.add_run(f'  {num}  {title}')
    run.font.size = Pt(14)
    p.add_run('.' * (50 - len(title)))
    doc.add_paragraph()  # spacing

doc.add_page_break()

# ============ 1. 绪论 ============
doc.add_heading('1  绪论', level=1)

doc.add_heading('1.1  项目背景', level=2)
doc.add_paragraph(
    '随着移动互联网的普及，餐饮行业的数字化需求日益增长。传统的中餐/快餐门店面临排队时间长、'
    '点餐效率低、人工成本高等问题。本项目旨在开发一款微信点餐小程序，让顾客通过手机自助完成'
    '浏览菜单、下单、支付的全流程，提升门店运营效率和顾客用餐体验。'
)

doc.add_heading('1.2  项目目标', level=2)
goals = [
    '实现顾客通过微信小程序自助点餐',
    '集成微信支付，完成在线支付闭环',
    '提供清晰的订单状态跟踪功能',
    '减轻前台点餐压力，提升翻台率',
]
for g in goals:
    p = doc.add_paragraph(g, style='List Bullet')

doc.add_heading('1.3  技术选型', level=2)
add_table(doc,
    ['层面', '技术', '说明'],
    [
        ['前端', '微信小程序原生', '微信生态内无缝使用'],
        ['后端', 'Node.js + Express', '轻量高效，适合小程序后端'],
        ['数据库', 'MySQL', '关系型数据，结构化存储'],
        ['认证', 'JWT', '无状态认证'],
        ['支付', '微信支付', '原生集成'],
    ]
)

doc.add_page_break()

# ============ 2. 需求分析 ============
doc.add_heading('2  需求分析', level=1)

doc.add_heading('2.1  用户角色', level=2)
doc.add_paragraph('顾客：小程序使用者，可通过微信授权登录，进行点餐、支付、查看订单等操作。')

doc.add_heading('2.2  功能需求', level=2)

doc.add_heading('2.2.1  用户模块', level=3)
add_table(doc,
    ['功能', '描述'],
    [
        ['微信授权登录', '使用 wx.login 获取 code，后端换取 openid，自动注册/登录'],
        ['用户信息展示', '展示昵称、头像等基本信息'],
    ]
)

doc.add_heading('2.2.2  菜品模块', level=3)
add_table(doc,
    ['功能', '描述'],
    [
        ['分类浏览', '按分类（热菜、凉菜、主食、汤品等）浏览菜品'],
        ['菜品列表', '展示菜品名称、图片、价格、月销量'],
        ['菜品详情', '查看菜品描述、规格选择'],
        ['菜品搜索', '按关键字搜索菜品'],
    ]
)

doc.add_heading('2.2.3  购物车模块', level=3)
add_table(doc,
    ['功能', '描述'],
    [
        ['添加菜品', '选择规格后加入购物车'],
        ['修改数量', '增减菜品数量'],
        ['删除菜品', '从购物车移除菜品'],
        ['查看购物车', '展示已选菜品、数量、总价'],
    ]
)

doc.add_heading('2.2.4  订单模块', level=3)
add_table(doc,
    ['功能', '描述'],
    [
        ['创建订单', '从购物车生成订单'],
        ['订单支付', '集成微信支付完成付款'],
        ['订单列表', '按状态查看订单（全部/待支付/已支付/已完成）'],
        ['订单详情', '查看订单包含的菜品、金额、状态'],
    ]
)

doc.add_heading('2.3  非功能需求', level=2)
nfrs = [
    '页面加载时间不超过 3 秒',
    '支持高峰期 100+ 用户同时在线',
    '遵循微信支付安全规范',
    '敏感信息（如 openid）加密存储',
]
for n in nfrs:
    p = doc.add_paragraph(n, style='List Bullet')

doc.add_heading('2.4  用例描述', level=2)
doc.add_paragraph('顾客通过微信小程序进行以下操作：')
use_cases = [
    '授权登录：打开小程序 → 微信授权 → 自动登录',
    '浏览菜品：查看分类 → 选择分类 → 浏览菜品列表 → 查看详情',
    '管理购物车：添加菜品 → 修改数量 → 删除菜品 → 查看购物车',
    '下单支付：提交订单 → 微信支付 → 支付成功',
    '查看订单：订单列表 → 订单详情 → 状态跟踪',
]
for uc in use_cases:
    p = doc.add_paragraph(uc, style='List Number')

doc.add_page_break()

# ============ 3. 总体设计 ============
doc.add_heading('3  总体设计', level=1)

doc.add_heading('3.1  系统架构', level=2)
doc.add_paragraph('系统采用前后端分离架构，前端为微信小程序，后端为 Node.js Express 应用，数据层使用 MySQL 数据库。')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[图 3-1  系统架构图]')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(11)

# ASCII architecture
arch_text = """
┌──────────────────────────────────┐
│       微信小程序（顾客端）          │
│  首页  详情页  购物车  订单列表    │
├──────────────────────────────────┤
│         HTTPS / JSON API          │
├──────────────────────────────────┤
│     Node.js Express 后端           │
│  路由层 → 控制层 → 数据访问层      │
├──────────────────────────────────┤
│            MySQL 数据库            │
└──────────────────────────────────┘
"""
p = doc.add_paragraph()
run = p.add_run(arch_text.strip())
run.font.name = 'Courier New'
run.font.size = Pt(10)

doc.add_heading('3.2  模块划分', level=2)
add_table(doc,
    ['模块', '职责', '关键接口'],
    [
        ['user', '微信登录、用户管理', '/api/user/*'],
        ['category', '菜品分类管理', '/api/categories'],
        ['dish', '菜品浏览、搜索', '/api/dishes/*'],
        ['order', '订单创建、状态管理', '/api/orders/*'],
        ['pay', '微信支付集成', '/api/pay/*'],
    ]
)

doc.add_heading('3.3  前后端交互流程', level=2)
doc.add_paragraph('小程序页面 → HTTP 请求 → Express 路由 → Controller → Model → MySQL → JSON 响应 → 小程序页面')

doc.add_page_break()

# ============ 4. 数据库设计 ============
doc.add_heading('4  数据库设计', level=1)

doc.add_heading('4.1  实体关系', level=2)
doc.add_paragraph('系统包含 5 个核心实体：User（用户）、Category（分类）、Dish（菜品）、Order（订单）、OrderItem（订单项）。')
doc.add_paragraph('')
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = p.add_run('[图 4-1  ER 关系图]')
run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
run.font.size = Pt(11)

er_text = """
User(1) ──── N Order(1) ──── N OrderItem
                                    │
Category(1) ── N Dish(1) ──────────┘
"""
p = doc.add_paragraph()
run = p.add_run(er_text.strip())
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_heading('4.2  数据表结构', level=2)

# user table
doc.add_heading('4.2.1  user（用户表）', level=3)
add_table(doc,
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PK AUTO_INCREMENT', '主键'],
        ['openid', 'VARCHAR(64)', 'UNIQUE NOT NULL', '微信 openid'],
        ['nickname', 'VARCHAR(50)', '', '昵称'],
        ['avatar_url', 'VARCHAR(255)', '', '头像 URL'],
        ['phone', 'VARCHAR(20)', '', '手机号（可选）'],
        ['create_time', 'DATETIME', 'DEFAULT NOW()', '创建时间'],
        ['update_time', 'DATETIME', 'ON UPDATE NOW()', '更新时间'],
    ]
)

# category table
doc.add_heading('4.2.2  category（菜品分类表）', level=3)
add_table(doc,
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PK AUTO_INCREMENT', '主键'],
        ['name', 'VARCHAR(30)', 'NOT NULL', '分类名称'],
        ['sort', 'INT', 'DEFAULT 0', '排序号'],
        ['create_time', 'DATETIME', 'DEFAULT NOW()', '创建时间'],
    ]
)

# dish table
doc.add_heading('4.2.3  dish（菜品表）', level=3)
add_table(doc,
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PK AUTO_INCREMENT', '主键'],
        ['category_id', 'INT', 'FK NOT NULL', '所属分类 ID'],
        ['name', 'VARCHAR(50)', 'NOT NULL', '菜品名称'],
        ['picture', 'VARCHAR(255)', '', '图片 URL'],
        ['price', 'DECIMAL(10,2)', 'NOT NULL', '价格'],
        ['description', 'VARCHAR(500)', '', '描述'],
        ['status', 'TINYINT', 'DEFAULT 1', '1=上架 0=下架'],
        ['sales', 'INT', 'DEFAULT 0', '月销量'],
        ['create_time', 'DATETIME', 'DEFAULT NOW()', '创建时间'],
    ]
)

# order table
doc.add_heading('4.2.4  order（订单表）', level=3)
add_table(doc,
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PK AUTO_INCREMENT', '主键'],
        ['order_no', 'VARCHAR(32)', 'UNIQUE NOT NULL', '订单号'],
        ['user_id', 'INT', 'FK NOT NULL', '用户 ID'],
        ['total_amount', 'DECIMAL(10,2)', 'NOT NULL', '订单总金额'],
        ['status', 'TINYINT', 'DEFAULT 0', '0=待支付 1=已支付 2=制作中 3=已完成 4=已取消'],
        ['remark', 'VARCHAR(200)', '', '备注'],
        ['pay_time', 'DATETIME', '', '支付时间'],
        ['create_time', 'DATETIME', 'DEFAULT NOW()', '创建时间'],
    ]
)

# order_item table
doc.add_heading('4.2.5  order_item（订单项表）', level=3)
add_table(doc,
    ['字段', '类型', '约束', '说明'],
    [
        ['id', 'INT', 'PK AUTO_INCREMENT', '主键'],
        ['order_id', 'INT', 'FK NOT NULL', '所属订单 ID'],
        ['dish_id', 'INT', 'FK NOT NULL', '菜品 ID'],
        ['dish_name', 'VARCHAR(50)', 'NOT NULL', '菜品名称（快照）'],
        ['dish_price', 'DECIMAL(10,2)', 'NOT NULL', '单价（快照）'],
        ['quantity', 'INT', 'DEFAULT 1', '数量'],
        ['subtotal', 'DECIMAL(10,2)', 'NOT NULL', '小计'],
    ]
)

doc.add_page_break()

# ============ 5. 接口设计 ============
doc.add_heading('5  接口设计', level=1)

doc.add_heading('5.1  通用说明', level=2)
p = doc.add_paragraph()
run = p.add_run('Base URL: ')
run.bold = True
p.add_run('https://api.example.com')

p = doc.add_paragraph()
run = p.add_run('Content-Type: ')
run.bold = True
p.add_run('application/json')

doc.add_paragraph('认证方式：除登录接口外，请求头携带 Authorization: Bearer <token>')
doc.add_paragraph('')

doc.add_paragraph('统一响应格式：')
resp_code = '''{
  "code": 0,
  "message": "success",
  "data": { }
}'''
p = doc.add_paragraph()
run = p.add_run(resp_code)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading('5.2  用户接口', level=2)

doc.add_heading('POST /api/user/login', level=3)
doc.add_paragraph('登录接口，使用微信临时 code 换取 openid 并返回 token。')
p = doc.add_paragraph()
run = p.add_run('Request:')
run.bold = True
p = doc.add_paragraph()
run = p.add_run('{ "code": "wx_login_code" }')
run.font.name = 'Courier New'
run.font.size = Pt(9)

p = doc.add_paragraph()
run = p.add_run('Response:')
run.bold = True
resp = '''{
  "code": 0,
  "data": {
    "token": "eyJhbGciOiJIUzI1NiIs...",
    "user": { "id": 1, "nickname": "微信用户", "avatar_url": "https://..." }
  }
}'''
p = doc.add_paragraph()
run = p.add_run(resp)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_heading('GET /api/user/info', level=3)
doc.add_paragraph('获取当前登录用户信息。')

doc.add_heading('5.3  菜品接口', level=2)

api_dishes = [
    ['GET', '/api/categories', '', '获取菜品分类列表'],
    ['GET', '/api/dishes', '?category_id=1', '获取菜品列表（按分类筛选）'],
    ['GET', '/api/dishes/:id', '', '获取菜品详情'],
    ['GET', '/api/dishes/search', '?keyword=鸡', '搜索菜品'],
]
add_table(doc, ['方法', '路径', '参数', '说明'], api_dishes)

doc.add_heading('5.4  订单接口', level=2)

api_orders = [
    ['POST', '/api/orders', '{ items, remark }', '创建订单'],
    ['GET', '/api/orders', '?status=0', '获取订单列表'],
    ['GET', '/api/orders/:id', '', '获取订单详情'],
]
add_table(doc, ['方法', '路径', '参数', '说明'], api_orders)

doc.add_heading('5.5  支付接口', level=2)

api_pay = [
    ['POST', '/api/pay/unified', '{ order_no }', '微信支付统一下单'],
    ['POST', '/api/pay/notify', '微信回调 XML', '支付结果回调通知'],
]
add_table(doc, ['方法', '路径', '参数', '说明'], api_pay)

doc.add_paragraph('')
doc.add_paragraph('创建订单示例：')
p = doc.add_paragraph()
run = p.add_run('''POST /api/orders
Request: { "items": [{"dish_id":1,"quantity":2}], "remark":"少辣" }
Response: { "code":0, "data": { "order_no":"20260515001", "total_amount":68.00 } }''')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_paragraph('')
doc.add_paragraph('支付示例：')
p = doc.add_paragraph()
run = p.add_run('''POST /api/pay/unified
Request: { "order_no": "20260515001" }
Response: {
  "code": 0,
  "data": {
    "payParams": {
      "nonceStr": "...", "package": "prepay_id=...",
      "paySign": "...", "signType": "MD5", "timeStamp": "..."
    }
  }
}''')
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_page_break()

# ============ 6. 详细设计 ============
doc.add_heading('6  详细设计', level=1)

doc.add_heading('6.1  页面结构', level=2)
add_table(doc,
    ['页面路径', '页面名称', '功能描述'],
    [
        ['pages/index/index', '首页', '分类导航 + 当前分类下的菜品列表'],
        ['pages/dish/detail', '菜品详情', '菜品大图、描述、价格、加入购物车'],
        ['pages/cart/cart', '购物车', '已选菜品列表、数量修改、结算'],
        ['pages/order/confirm', '确认订单', '订单项确认、备注填写、去支付'],
        ['pages/order/list', '订单列表', '按状态 tab 切换展示订单'],
        ['pages/order/detail', '订单详情', '订单信息、菜品清单、状态跟踪'],
        ['pages/user/user', '个人中心', '用户信息、订单入口'],
    ]
)

doc.add_heading('6.2  购物车设计', level=2)
doc.add_paragraph('使用微信小程序本地缓存 wx.setStorageSync 存储购物车数据，每次增删改操作后同步更新。')
doc.add_paragraph('数据结构：')
p = doc.add_paragraph()
run = p.add_run('''{
  "items": [
    { "dish_id": 1, "name": "宫保鸡丁", "price": 32.00, "quantity": 2 },
    { "dish_id": 3, "name": "米饭", "price": 3.00, "quantity": 1 }
  ],
  "totalCount": 3,
  "totalAmount": 67.00
}''')
run.font.name = 'Courier New'
run.font.size = Pt(9)
doc.add_paragraph('下单成功后清空购物车。')

doc.add_heading('6.3  订单状态机', level=2)
p = doc.add_paragraph()
run = p.add_run('0(待支付) → 1(已支付) → 2(制作中) → 3(已完成)')
run.font.name = 'Courier New'
run.font.size = Pt(11)
p = doc.add_paragraph()
run = p.add_run('0(待支付) → 4(已取消)')
run.font.name = 'Courier New'
run.font.size = Pt(11)

doc.add_heading('6.4  微信登录流程', level=2)
login_steps = [
    '小程序 wx.login() 获取临时 code',
    '请求 POST /api/user/login { code }',
    '后端向微信服务器 jscode2session 接口换取 openid',
    '根据 openid 查询/创建用户记录',
    '生成 JWT token（payload 包含 user_id, openid）',
    '返回 token + 用户信息',
    '后续请求通过 Authorization 头携带 token',
    'auth 中间件解析 token 并设置 req.user',
]
for s in login_steps:
    p = doc.add_paragraph(s, style='List Number')

doc.add_heading('6.5  微信支付流程', level=2)
pay_steps = [
    '用户提交订单后请求 POST /api/pay/unified',
    '后端调微信支付统一下单 API，传入 appid, mch_id, openid, total_fee 等参数',
    '微信返回 prepay_id',
    '后端组装小程序支付参数（timeStamp, nonceStr, package, paySign）',
    '返回 payParams 给小程序端',
    '小程序端调用 wx.requestPayment(payParams)',
    '用户确认支付，微信异步回调 notify_url',
    '后端接收回调，验证签名，更新订单状态为已支付',
]
for s in pay_steps:
    p = doc.add_paragraph(s, style='List Number')

doc.add_heading('6.6  项目目录结构', level=2)
dir_text = '''点餐小程序/
├── client/                     # 小程序前端
│   ├── app.js / app.json / app.wxss
│   ├── pages/
│   │   ├── index/              # 首页
│   │   ├── dish/detail/        # 菜品详情
│   │   ├── cart/               # 购物车
│   │   ├── order/              # 订单 (confirm/list/detail)
│   │   └── user/               # 个人中心
│   └── utils/api.js            # HTTP 请求封装
│
├── server/                     # Node.js 后端
│   ├── app.js                  # 入口文件
│   ├── package.json
│   ├── config/index.js         # 配置
│   ├── routes/                 # 路由
│   ├── controllers/            # 控制器
│   ├── models/                 # 数据模型
│   ├── middleware/             # 中间件
│   └── utils/                  # 工具函数
│
└── docs/SQL/init.sql           # 数据库初始化脚本'''
p = doc.add_paragraph()
run = p.add_run(dir_text)
run.font.name = 'Courier New'
run.font.size = Pt(9)

doc.add_page_break()

# ============ 7. API 汇总 ============
doc.add_heading('7  API 接口汇总', level=1)

all_apis = [
    ['POST', '/api/user/login', '否', '微信登录'],
    ['GET', '/api/user/info', '是', '获取用户信息'],
    ['GET', '/api/categories', '是', '获取分类列表'],
    ['GET', '/api/dishes', '是', '获取菜品列表'],
    ['GET', '/api/dishes/:id', '是', '获取菜品详情'],
    ['GET', '/api/dishes/search', '是', '搜索菜品'],
    ['POST', '/api/orders', '是', '创建订单'],
    ['GET', '/api/orders', '是', '获取订单列表'],
    ['GET', '/api/orders/:id', '是', '获取订单详情'],
    ['POST', '/api/pay/unified', '是', '微信统一下单'],
    ['POST', '/api/pay/notify', '否', '支付回调'],
]
add_table(doc, ['方法', '路径', '需登录', '说明'], all_apis)

# ============ 保存文件 ============
output_path = 'D:\\互助养老3\\docs\\superpowers\\specs\\点餐小程序项目设计文档.docx'
doc.save(output_path)
print(f'文档已生成: {output_path}')
